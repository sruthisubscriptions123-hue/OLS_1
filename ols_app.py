"""
OLS Regression Analysis Tool  —  Streamlit App  (v2)
=====================================================
Run:  streamlit run ols_app.py

Requirements:
    pip install streamlit pandas numpy scipy statsmodels patsy plotly openpyxl python-docx kaleido
"""

import io
import warnings
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import streamlit as st
from scipy import stats
from patsy import dmatrices, EvalEnvironment
import statsmodels.formula.api as smf
from statsmodels.stats.stattools import durbin_watson
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="OLS Regression Tool",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  LIGHT THEME CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
/* Hide default sidebar toggle arrow on small screens */
[data-testid="collapsedControl"] { display: none; }

/* Main content area */
.main .block-container {
    padding: 1.8rem 2.5rem 4rem 2.5rem;
    max-width: 1050px;
}

/* Typography */
h1 { font-size: 1.7rem !important; font-weight: 700 !important; color: #1a2340 !important; }
h2 { font-size: 1.25rem !important; font-weight: 700 !important; color: #1a3a6b !important;
     border-left: 4px solid #1a6bb5; padding-left: 12px; margin-top: 1.4rem !important; }
h3 { font-size: 1.05rem !important; font-weight: 600 !important; color: #2c4a7c !important; }
h4 { font-size: 0.95rem !important; font-weight: 600 !important; color: #3a5a8a !important; }
p, li { color: #2d3748 !important; font-size: 0.95rem; }
label { color: #374151 !important; font-size: 0.88rem !important; font-weight: 500 !important; }

/* Step badge */
.step-badge {
    display: inline-block; padding: 4px 14px; border-radius: 20px; font-size: 0.8rem;
    font-weight: 700; margin-bottom: 0.7rem; letter-spacing: 0.5px;
    background: #e8f0fe; color: #1a6bb5; border: 1px solid #b3d0f5;
}

/* Sidebar nav link styling */
.nav-active {
    background: #1a6bb5 !important; color: #fff !important;
    border-radius: 7px; padding: 8px 14px; font-weight: 700;
    display: block; margin-bottom: 4px;
}
.nav-inactive {
    background: transparent; color: #374151;
    border-radius: 7px; padding: 8px 14px; font-weight: 500;
    display: block; margin-bottom: 4px;
}
.nav-done {
    background: #eaf7f0; color: #2d7a5f;
    border-radius: 7px; padding: 8px 14px; font-weight: 600;
    display: block; margin-bottom: 4px;
}

/* Verdict colors */
.verdict-ok   { color: #1a7a3a; font-weight: 700; }
.verdict-warn { color: #b35c00; font-weight: 700; }

/* VIF badge */
.vif-ok   { background:#d4edda; color:#155724; border-radius:5px; padding:2px 9px; font-size:0.85rem; font-weight:600; }
.vif-mod  { background:#fff3cd; color:#856404; border-radius:5px; padding:2px 9px; font-size:0.85rem; font-weight:600; }
.vif-high { background:#f8d7da; color:#721c24; border-radius:5px; padding:2px 9px; font-size:0.85rem; font-weight:600; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────
STEPS = ["Load Data", "Variable Types", "Model Build",
         "Coefficients", "Fit & Assumptions", "Diagnostics", "Multicollinearity"]

defaults = {
    "current_step": 1,
    "raw_df": None,
    "confirmed_df": None,
    "col_types": {},
    "dv": None,
    "selected_ivs": [],
    "centering": {},
    "alpha": 0.05,
    "decimals": 3,
    "ci_level": 0.95,
    "patsy_formula": "",
    "model_result": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ─────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────
def go_step(n):
    st.session_state["current_step"] = n


def fmt(v, d=None):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    d = d if d is not None else st.session_state["decimals"]
    return f"{v:.{d}f}"


def fmt_p(v):
    if v < 0.001:
        return "<.001"
    return f"{v:.3f}"


def badge(label):
    return f'<span class="step-badge">{label}</span>'


def infer_type(series: pd.Series) -> str:
    if pd.api.types.is_datetime64_any_dtype(series):
        return "datetime"
    try:
        pd.to_numeric(series.dropna(), errors="raise")
        return "numeric"
    except Exception:
        pass
    n_uniq = series.nunique()
    n_total = series.count()
    if n_uniq <= 15 or (n_total > 0 and n_uniq / n_total <= 0.30):
        return "categorical"
    return "object"


def apply_confirmed_types(df: pd.DataFrame, col_types: dict) -> pd.DataFrame:
    df = df.copy()
    for col, ctype in col_types.items():
        if col not in df.columns:
            continue
        if ctype == "numeric":
            df[col] = pd.to_numeric(df[col], errors="coerce")
        elif ctype == "categorical":
            df[col] = df[col].astype(str).astype("category")
        elif ctype == "datetime":
            df[col] = pd.to_datetime(df[col], errors="coerce")
        else:
            df[col] = df[col].astype(str)
    return df


def build_patsy_formula() -> str:
    dv  = st.session_state["dv"]
    ivs = st.session_state["selected_ivs"]
    if not dv or not ivs:
        return ""
    terms = []
    for col in ivs:
        t = st.session_state["col_types"].get(col, "numeric")
        terms.append(f"C({col})" if t == "categorical" else col)
    return f"{dv} ~ {' + '.join(terms)}"


# ── Light Plotly theme ────────────────────────
PLOTLY_LIGHT = dict(
    paper_bgcolor="#ffffff",
    plot_bgcolor="#f8f9fc",
    font=dict(color="#2d3748", family="Arial, sans-serif", size=13),
    xaxis=dict(gridcolor="#e2e8f0", linecolor="#cbd5e0", zerolinecolor="#cbd5e0",
               tickfont=dict(size=12), title_font=dict(size=13)),
    yaxis=dict(gridcolor="#e2e8f0", linecolor="#cbd5e0", zerolinecolor="#cbd5e0",
               tickfont=dict(size=12), title_font=dict(size=13)),
    margin=dict(t=55, b=55, l=70, r=30),
    legend=dict(bgcolor="#ffffff", font=dict(color="#2d3748", size=12),
                bordercolor="#e2e8f0", borderwidth=1),
)


# Separate layout dict for heatmaps — no xaxis/yaxis keys to avoid conflicts
PLOTLY_LIGHT_HEATMAP = dict(
    paper_bgcolor="#ffffff",
    plot_bgcolor="#f8f9fc",
    font=dict(color="#2d3748", family="Arial, sans-serif", size=13),
    margin=dict(t=55, b=55, l=70, r=30),
)
def _plot_title(text):
    return dict(text=text, font=dict(size=15, color="#1a2340"))

# ─────────────────────────────────────────────
#  OLS ENGINE
# ─────────────────────────────────────────────
def _run_ols(df, formula, alpha, ci_level, centering):
    df = df.copy()
    for col, do_center in centering.items():
        if do_center and col in df.columns:
            df[col] = df[col] - df[col].mean()
    df = df.dropna()

    model = smf.ols(formula=formula, data=df)
    fit   = model.fit()

    ci     = fit.conf_int(alpha=1 - ci_level)
    params = fit.params
    pvals  = fit.pvalues
    bse    = fit.bse
    tvals  = fit.tvalues

    coef_rows = []
    for name in params.index:
        coef_rows.append({
            "term":  str(name),
            "coef":  float(params[name]),
            "se":    float(bse[name]),
            "tstat": float(tvals[name]),
            "pval":  float(pvals[name]),
            "ci_lo": float(ci.loc[name, 0]),
            "ci_hi": float(ci.loc[name, 1]),
            "sig":   bool(pvals[name] < alpha),
        })

    influence   = fit.get_influence()
    resid       = fit.resid.values
    fitted_vals = fit.fittedvalues.values
    leverage    = influence.hat_matrix_diag
    cooks_d     = influence.cooks_distance[0]
    std_resid   = influence.resid_studentized_internal

    sw_sample = resid[:5000] if len(resid) > 5000 else resid
    sw_stat, sw_p = stats.shapiro(sw_sample)
    dw_val = float(durbin_watson(resid))
    bp_lm, bp_p, _, _ = het_breuschpagan(resid, fit.model.exog)
    lin_r, lin_p = stats.pearsonr(fitted_vals, resid)

    # VIF — only for numeric IVs, requires ≥2 numeric predictors
    vif_rows = []
    try:
        exog = fit.model.exog
        exog_df = pd.DataFrame(exog, columns=fit.model.exog_names)
        # Drop Intercept column for VIF
        num_cols = [c for c in exog_df.columns if c != "Intercept"]
        if len(num_cols) >= 2:
            for j, col_name in enumerate(exog_df.columns):
                if col_name == "Intercept":
                    continue
                try:
                    vif_val = variance_inflation_factor(exog_df.values, j)
                    vif_rows.append({"term": col_name, "vif": float(vif_val)})
                except Exception:
                    vif_rows.append({"term": col_name, "vif": float("nan")})
    except Exception:
        vif_rows = []

    # Correlation matrix of numeric predictors
    ivs_list = st.session_state["selected_ivs"]
    col_types = st.session_state["col_types"]
    num_ivs = [v for v in ivs_list if col_types.get(v) == "numeric"]
    corr_matrix = None
    if len(num_ivs) >= 2:
        corr_matrix = df[num_ivs].corr().round(3).to_dict()

    return {
        "coef_rows":   coef_rows,
        "rsq":         float(fit.rsquared),
        "rsq_adj":     float(fit.rsquared_adj),
        "fstat":       float(fit.fvalue),
        "fpval":       float(fit.f_pvalue),
        "df_model":    int(fit.df_model),
        "df_resid":    int(fit.df_resid),
        "aic":         float(fit.aic),
        "bic":         float(fit.bic),
        "rmse":        float(np.sqrt(fit.mse_resid)),
        "n":           int(fit.nobs),
        "sw_stat":     float(sw_stat), "sw_p":  float(sw_p),
        "dw":          dw_val,
        "bp_stat":     float(bp_lm),   "bp_p":  float(bp_p),
        "lin_r":       float(lin_r),   "lin_p": float(lin_p),
        "fitted":      fitted_vals.tolist(),
        "resid":       resid.tolist(),
        "leverage":    leverage.tolist(),
        "cooks_d":     cooks_d.tolist(),
        "std_resid":   std_resid.tolist(),
        "obs_idx":     list(range(len(resid))),
        "formula":     formula,
        "dv":          formula.split("~")[0].strip(),
        "alpha":       alpha,
        "ci_level":    ci_level,
        "vif_rows":    vif_rows,
        "corr_matrix": corr_matrix,
        "num_ivs":     num_ivs,
    }


# ─────────────────────────────────────────────
#  SIDEBAR NAVIGATION
# ─────────────────────────────────────────────
cs = st.session_state["current_step"]

with st.sidebar:
    st.markdown("""
    <div style="padding:10px 0 18px 0;">
      <div style="font-size:1.15rem;font-weight:800;color:#1a2340;">📊 OLS Regression</div>
      <div style="font-size:0.8rem;color:#6b7280;margin-top:2px;">Analysis Tool</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("<div style='font-size:0.75rem;font-weight:700;color:#9ca3af;"
                "text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;'>"
                "Steps</div>", unsafe_allow_html=True)

    for i, step_name in enumerate(STEPS):
        step_num = i + 1
        is_current = step_num == cs
        is_done    = step_num < cs

        if is_current:
            cls = "nav-active"
            prefix = f"▶  {step_num}. "
        elif is_done:
            cls = "nav-done"
            prefix = f"✓  {step_num}. "
        else:
            cls = "nav-inactive"
            prefix = f"    {step_num}. "

        # Only allow navigation to completed/current steps
        if is_done or is_current:
            if st.button(f"{prefix}{step_name}", key=f"nav_{step_num}",
                         use_container_width=True):
                go_step(step_num)
                st.rerun()
        else:
            st.markdown(
                f"<div style='color:#9ca3af;padding:8px 14px;font-size:0.9rem;"
                f"margin-bottom:4px;'>{prefix}{step_name}</div>",
                unsafe_allow_html=True
            )

    st.markdown("---")
    if st.session_state["model_result"]:
        r_s = st.session_state["model_result"]
        st.markdown("<div style='font-size:0.75rem;font-weight:700;color:#9ca3af;"
                    "text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;'>"
                    "Quick Stats</div>", unsafe_allow_html=True)
        st.metric("R²",      f"{r_s['rsq']:.4f}")
        st.metric("Adj. R²", f"{r_s['rsq_adj']:.4f}")
        st.metric("N",       f"{r_s['n']:,}")


# ─────────────────────────────────────────────
#  MAIN CONTENT
# ─────────────────────────────────────────────

# ═══════════════════════════════════════════════════════════════════════
#  STEP 1 — LOAD DATA
# ═══════════════════════════════════════════════════════════════════════
if cs == 1:
    st.markdown(badge("Step 1"), unsafe_allow_html=True)
    st.markdown("## Load Data")

    uploaded = st.file_uploader(
        "Upload your Excel file (.xlsx / .xls)", type=["xlsx", "xls"],
        help="Maximum 10,000 rows recommended."
    )

    if uploaded:
        try:
            xl = pd.ExcelFile(uploaded)
            sheet_names = xl.sheet_names

            c1, c2 = st.columns(2)
            with c1:
                sheet = st.selectbox("Sheet", sheet_names)
            with c2:
                header_row = st.selectbox("Header row", [0, 1, 2],
                                          format_func=lambda x: f"Row {x+1}")

            df_raw = xl.parse(sheet, header=header_row)
            df_raw = df_raw.loc[:, df_raw.columns.notna()]
            df_raw.columns = [str(c).strip() for c in df_raw.columns]

            st.markdown("### Preview (first 5 rows)")
            st.dataframe(df_raw.head(), use_container_width=True, hide_index=True)

            st.markdown("### Dataset dimensions")
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Rows",         f"{len(df_raw):,}")
            m2.metric("Columns",      len(df_raw.columns))
            m3.metric("Missing cells", int(df_raw.isnull().sum().sum()))
            m4.metric("Sheet",        sheet)

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("✓ Confirm & proceed to variable types", type="primary"):
                st.session_state["raw_df"] = df_raw
                go_step(2)
                st.rerun()

        except Exception as e:
            st.error(f"Could not read file: {e}")
    else:
        st.info("👆 Upload an Excel file to begin.")


# ═══════════════════════════════════════════════════════════════════════
#  STEP 2 — VARIABLE TYPES
# ═══════════════════════════════════════════════════════════════════════
elif cs == 2:
    st.markdown(badge("Step 2"), unsafe_allow_html=True)
    st.markdown("## Variable Types")
    st.info(
        "Review the inferred data types below. Adjust any column before continuing — "
        "all downstream steps will use the types confirmed here."
    )

    df = st.session_state["raw_df"]
    TYPE_OPTIONS = ["numeric", "categorical", "object", "datetime"]
    TYPE_COLORS  = {
        "numeric":     "#1a6bb5",
        "categorical": "#d97706",
        "object":      "#059669",
        "datetime":    "#7c3aed",
    }

    new_types = {}
    header_cols = st.columns([0.5, 2.5, 3, 1.5, 2])
    for h, c in zip(["#", "Variable", "Sample values", "Inferred", "Change type"], header_cols):
        c.markdown(f"<span style='color:#6b7280;font-size:0.8rem;font-weight:700;"
                   f"text-transform:uppercase;letter-spacing:.7px'>{h}</span>",
                   unsafe_allow_html=True)
    st.markdown("<hr style='margin:4px 0 10px'>", unsafe_allow_html=True)

    for i, col in enumerate(df.columns):
        inferred  = infer_type(df[col])
        prev_type = st.session_state["col_types"].get(col, inferred)

        sample_vals = df[col].dropna().astype(str).head(3).tolist()
        sample_str  = ",  ".join(sample_vals) if sample_vals else "—"

        c0, c1, c2, c3, c4 = st.columns([0.5, 2.5, 3, 1.5, 2])
        c0.markdown(f"<span style='color:#9ca3af;font-size:0.85rem'>{i+1}</span>",
                    unsafe_allow_html=True)
        c1.markdown(f"<code style='font-size:0.88rem'>{col}</code>", unsafe_allow_html=True)
        c2.markdown(
            f"<span style='color:#6b7280;font-size:0.85rem'>{sample_str[:60]}</span>",
            unsafe_allow_html=True
        )
        color = TYPE_COLORS.get(inferred, "#6b7280")
        c3.markdown(
            f"<span style='background:#f3f4f6;color:{color};"
            f"border:1px solid {color};border-radius:5px;padding:2px 9px;font-size:0.82rem;"
            f"font-family:monospace;font-weight:600'>{inferred}</span>",
            unsafe_allow_html=True
        )
        chosen = c4.selectbox(
            f"__{col}__", TYPE_OPTIONS,
            index=TYPE_OPTIONS.index(prev_type),
            key=f"type_{col}", label_visibility="collapsed"
        )
        new_types[col] = chosen

    st.markdown("<br>", unsafe_allow_html=True)
    bc1, bc2 = st.columns([1, 5])
    with bc1:
        if st.button("‹ Back"):
            go_step(1); st.rerun()
    with bc2:
        if st.button("✓ Confirm types & build model", type="primary"):
            st.session_state["col_types"] = new_types
            confirmed = apply_confirmed_types(st.session_state["raw_df"], new_types)
            st.session_state["confirmed_df"] = confirmed
            go_step(3); st.rerun()


# ═══════════════════════════════════════════════════════════════════════
#  STEP 3 — MODEL BUILD  (interactions removed)
# ═══════════════════════════════════════════════════════════════════════
elif cs == 3:
    st.markdown(badge("Step 3"), unsafe_allow_html=True)
    st.markdown("## Model Specification")

    df        = st.session_state["confirmed_df"]
    cols      = list(df.columns)
    col_types = st.session_state["col_types"]

    st.markdown("### Dependent variable & settings")
    r1c1, r1c2, r1c3, r1c4 = st.columns(4)

    with r1c1:
        dv_options = ["— select —"] + cols
        prev_dv    = st.session_state["dv"] or "— select —"
        dv_idx     = dv_options.index(prev_dv) if prev_dv in dv_options else 0
        dv         = st.selectbox("Dependent variable (response)", dv_options, index=dv_idx)
        dv         = None if dv == "— select —" else dv
        st.session_state["dv"] = dv

    with r1c2:
        alpha = st.selectbox("Significance level (α)",
                             [0.05, 0.01, 0.10],
                             format_func=lambda x: str(x),
                             index=[0.05, 0.01, 0.10].index(st.session_state["alpha"]))
        st.session_state["alpha"] = alpha

    with r1c3:
        ci_level = st.selectbox("CI level",
                                [0.95, 0.99, 0.90],
                                format_func=lambda x: f"{int(x*100)}%",
                                index=[0.95, 0.99, 0.90].index(st.session_state["ci_level"]))
        st.session_state["ci_level"] = ci_level

    with r1c4:
        decimals = st.selectbox("Decimal places", [2, 3, 4],
                                index=[2, 3, 4].index(st.session_state["decimals"]))
        st.session_state["decimals"] = decimals

    st.markdown("---")

    st.markdown("### Independent variables (predictors)")
    st.caption("Tick each variable you want to include as a predictor.")
    avail_ivs = [c for c in cols if c != dv]
    prev_ivs  = [v for v in st.session_state["selected_ivs"] if v in avail_ivs]

    # Render checkboxes in a 4-column grid — stable, no dropdown-jump
    TYPE_CHIP = {
        "numeric":     ("#1a6bb5", "#e8f0fe"),
        "categorical": ("#92400e", "#fef3c7"),
        "object":      ("#065f46", "#d1fae5"),
        "datetime":    ("#5b21b6", "#ede9fe"),
    }
    n_cols = 4
    grid_cols = st.columns(n_cols)
    selected_ivs = []
    for idx, iv in enumerate(avail_ivs):
        t    = col_types.get(iv, "numeric")
        fg, bg = TYPE_CHIP.get(t, ("#374151", "#f3f4f6"))
        with grid_cols[idx % n_cols]:
            # Tiny type badge above each checkbox
            st.markdown(
                f"<span style='font-size:0.72rem;font-weight:700;color:{fg};"
                f"background:{bg};border-radius:4px;padding:1px 7px;"
                f"display:inline-block;margin-bottom:2px'>{t}</span>",
                unsafe_allow_html=True,
            )
            checked = st.checkbox(
                iv,
                value=(iv in prev_ivs),
                key=f"iv_cb_{iv}",
            )
            if checked:
                selected_ivs.append(iv)

    st.session_state["selected_ivs"] = selected_ivs

    if selected_ivs:
        st.caption(f"**{len(selected_ivs)} predictor(s) selected:** " +
                   " · ".join(f"{iv} ({col_types.get(iv,'numeric')})" for iv in selected_ivs))
    else:
        st.caption("No predictors selected yet.")

    st.markdown("---")

    # Mean-centering
    num_selected = [iv for iv in selected_ivs if col_types.get(iv) == "numeric"]
    if num_selected:
        st.markdown("### Mean-centering *(optional — numeric predictors only)*")
        center_cols = st.columns(min(len(num_selected), 4))
        for j, col_name in enumerate(num_selected):
            with center_cols[j % 4]:
                prev = st.session_state["centering"].get(col_name, False)
                val  = st.checkbox(col_name, value=prev, key=f"center_{col_name}")
                st.session_state["centering"][col_name] = val
        st.caption("Checked variables will be mean-centered before model fitting.")
        st.markdown("---")

    # Formula preview
    formula = build_patsy_formula()
    st.session_state["patsy_formula"] = formula
    st.markdown("### Patsy formula preview")
    if formula:
        st.code(formula, language="r")
    else:
        st.caption("Select a dependent variable and at least one predictor to preview.")

    st.markdown("<br>", unsafe_allow_html=True)
    bc1, bc2 = st.columns([1, 5])
    with bc1:
        if st.button("‹ Back"):
            go_step(2); st.rerun()
    with bc2:
        run_clicked = st.button("▶ Run OLS model", type="primary",
                                disabled=(not formula))

    if run_clicked:
        with st.spinner("Fitting OLS model…"):
            try:
                result = _run_ols(
                    df=st.session_state["confirmed_df"],
                    formula=formula,
                    alpha=alpha,
                    ci_level=ci_level,
                    centering=st.session_state["centering"],
                )
                st.session_state["model_result"] = result
                go_step(4); st.rerun()
            except Exception as e:
                st.error(f"Model error: {e}")


# ═══════════════════════════════════════════════════════════════════════
#  STEP 4 — COEFFICIENT TABLE
# ═══════════════════════════════════════════════════════════════════════
elif cs == 4:
    st.markdown(badge("Step 4"), unsafe_allow_html=True)
    st.markdown("## Model Output — Coefficients")

    r = st.session_state["model_result"]
    if r is None:
        st.warning("No model results. Please run the model first.")
        if st.button("‹ Go back to Model Build"):
            go_step(3); st.rerun()
        st.stop()

    d     = st.session_state["decimals"]
    alpha = r["alpha"]
    ci    = int(r["ci_level"] * 100)

    # Model spec
    with st.expander("📋 Model specification", expanded=True):
        spec_rows = [
            ("Dependent variable",      r["dv"]),
            ("Independent variables",   ", ".join(st.session_state["selected_ivs"])),
            ("Patsy formula",           r["formula"]),
            ("Significance level (α)",  str(alpha)),
            ("CI level",                f"{ci}%"),
            ("Mean-centered",           ", ".join([k for k, v in st.session_state["centering"].items() if v]) or "none"),
            ("Observations used",       str(r["n"])),
        ]
        for k, v in spec_rows:
            c1, c2 = st.columns([2, 5])
            c1.markdown(f"<span style='color:#6b7280;font-size:0.92rem;font-weight:600'>{k}</span>",
                        unsafe_allow_html=True)
            c2.markdown(f"<code style='font-size:0.9rem'>{v}</code>",
                        unsafe_allow_html=True)

    st.markdown("---")

    # Coefficient table
    ci_lo_label = f"CI {(1-r['ci_level'])/2*100:.1f}%"
    ci_hi_label = f"CI {(1-(1-r['ci_level'])/2)*100:.1f}%"

    rows = []
    for row in r["coef_rows"]:
        rows.append({
            "Term":        row["term"],
            "Coefficient": f"{row['coef']:.{d}f}",
            "Std. Error":  f"{row['se']:.{d}f}",
            "t-statistic": f"{row['tstat']:.{d}f}",
            "p-value":     fmt_p(row["pval"]),
            ci_lo_label:   f"{row['ci_lo']:.{d}f}",
            ci_hi_label:   f"{row['ci_hi']:.{d}f}",
            "Significant": "✅ Yes" if row["sig"] else "❌ No",
        })

    coef_df = pd.DataFrame(rows)

    def highlight_row(row):
        if "Yes" in str(row["Significant"]):
            return ["background-color: #d4edda"] * len(row)
        elif "No" in str(row["Significant"]) and row["Term"] != "Intercept":
            return ["background-color: #fde8e8"] * len(row)
        return [""] * len(row)

    styled = (
        coef_df.style
        .apply(highlight_row, axis=1)
        .set_properties(**{"font-size": "14px", "font-family": "Arial, sans-serif",
                           "text-align": "right"})
        .set_properties(subset=["Term", "Significant"], **{"text-align": "left"})
        .set_table_styles([
            {"selector": "th", "props": [
                ("background-color", "#1a3a6b"),
                ("color", "#ffffff"),
                ("font-size", "13px"),
                ("font-weight", "700"),
                ("text-align", "center"),
                ("padding", "10px 12px"),
            ]},
            {"selector": "td", "props": [
                ("font-size", "14px"),
                ("padding", "9px 12px"),
            ]},
        ])
    )
    st.dataframe(styled, use_container_width=True, hide_index=True, height=400)

    st.markdown(
        "🟢 **Green rows** = p < α (statistically significant)&nbsp;&nbsp;&nbsp;"
        "🔴 **Red rows** = p ≥ α (not significant)",
    )

    # F-test summary
    st.markdown("---")
    st.markdown("### Overall Model Fit — F-test")

    f_sig = r["fpval"] < alpha
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("F-statistic",  f"{r['fstat']:.3f}",
              help=f"df({r['df_model']}, {r['df_resid']})")
    m2.metric("p-value (F)",  fmt_p(r["fpval"]),
              delta="Significant" if f_sig else "Not significant",
              delta_color="normal" if f_sig else "inverse")
    m3.metric("R²",           f"{r['rsq']:.4f}")
    m4.metric("Adj. R²",      f"{r['rsq_adj']:.4f}")
    m5.metric("RMSE",         f"{r['rmse']:.{d}f}")
    m6.metric("AIC / BIC",    f"{r['aic']:.1f} / {r['bic']:.1f}")

    st.markdown("---")
    st.markdown(
        "**How to read the coefficient table:** The **Coefficient** shows the expected change in "
        f"{r['dv']} for a one-unit increase in each predictor (holding others constant). "
        "**Std. Error** measures precision. The **t-statistic** tests whether the coefficient "
        f"differs from zero. **p-value** < {alpha} indicates statistical significance. "
        "The confidence interval gives the plausible range for the true coefficient."
    )

    st.markdown("<br>", unsafe_allow_html=True)
    bc1, bc2 = st.columns([1, 5])
    with bc1:
        if st.button("‹ Modify model"):
            go_step(3); st.rerun()
    with bc2:
        if st.button("Fit & Assumptions ›", type="primary"):
            go_step(5); st.rerun()


# ═══════════════════════════════════════════════════════════════════════
#  STEP 5 — FIT METRICS & ASSUMPTIONS
# ═══════════════════════════════════════════════════════════════════════
elif cs == 5:
    st.markdown(badge("Step 5"), unsafe_allow_html=True)
    st.markdown("## Fit Metrics & Assumption Tests")

    r     = st.session_state["model_result"]
    alpha = r["alpha"]
    d     = st.session_state["decimals"]

    if r is None:
        st.warning("Run the model first.")
        if st.button("‹ Back"):
            go_step(3); st.rerun()
        st.stop()

    # Overall fit metrics
    st.markdown("### Overall Fit")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("R²",      f"{r['rsq']:.4f}",
              help=f"Model explains {r['rsq']*100:.1f}% of variance in {r['dv']}.")
    m2.metric("Adj. R²", f"{r['rsq_adj']:.4f}",
              help="Penalises for extra predictors. Use for model comparison.")
    m3.metric("RMSE",    f"{r['rmse']:.{d}f}",
              help=f"Average prediction error in units of {r['dv']}.")
    m4.metric("N",       f"{r['n']:,}",
              help="Rows after listwise deletion of missing values.")
    m5.metric("AIC",     f"{r['aic']:.1f}",
              help="Lower is better when comparing models on the same data.")
    m6.metric("BIC",     f"{r['bic']:.1f}",
              help="Stricter than AIC; penalises complexity more heavily.")

    st.markdown("---")
    st.markdown("### Assumption Tests")

    def verdict_md(ok):
        return ("✅ **Supported**" if ok else "⚠️ **Caution — review needed**")

    # ── 1. Linearity ──────────────────────────────────────
    with st.expander("📐 Test 1: Linearity  —  Residuals vs Fitted", expanded=True):
        lin_ok = abs(r["lin_r"]) < 0.1
        st.markdown(
            f"**Pearson r (residuals ~ fitted values):** `{r['lin_r']:.4f}`,  "
            f"p = `{fmt_p(r['lin_p'])}`  →  {verdict_md(lin_ok)}"
        )
        fig1 = go.Figure()
        fig1.add_trace(go.Scatter(
            x=r["fitted"], y=r["resid"], mode="markers",
            marker=dict(color="#1a6bb5", opacity=0.6, size=7),
            name="Residuals"
        ))
        fig1.add_hline(y=0, line_dash="dash", line_color="#e53e3e", line_width=2)
        fig1.update_layout(**PLOTLY_LIGHT,
                           xaxis_title=_plot_title("Fitted Values"),
                           yaxis_title="Residuals",
                           title="Residuals vs Fitted Values")
        st.plotly_chart(fig1, use_container_width=True)

        st.markdown("**📖 Interpretation:**")
        if lin_ok:
            st.success(
                "The residuals show no systematic linear pattern with fitted values (r ≈ 0). "
                "This supports the linearity assumption — your model captures the relationship "
                "without a clear trend in the errors. Points should scatter randomly around the "
                "horizontal zero line."
            )
        else:
            st.warning(
                f"Residuals correlate with fitted values (r = {r['lin_r']:.3f}), suggesting a "
                "non-linear relationship that your current model does not capture. Consider: "
                "(1) adding polynomial terms, (2) applying log/sqrt transformations to skewed "
                "predictors, or (3) checking for omitted variables."
            )

    # ── 2. Independence (Durbin-Watson) ──────────────────
    with st.expander("🔄 Test 2: Independence  —  Durbin-Watson", expanded=True):
        dw_ok = 1.5 < r["dw"] < 2.5
        st.markdown(
            f"**Durbin-Watson statistic:** `{r['dw']:.4f}`  →  {verdict_md(dw_ok)}"
        )
        fig2 = go.Figure()
        fig2.add_trace(go.Scatter(
            x=list(range(1, len(r["resid"])+1)),
            y=r["resid"],
            mode="lines+markers",
            line=dict(color="#1a6bb5", width=1.8),
            marker=dict(color="#1a6bb5", size=5, opacity=0.7),
            name="Residuals"
        ))
        fig2.add_hline(y=0, line_dash="dash", line_color="#e53e3e", line_width=2)
        fig2.update_layout(**PLOTLY_LIGHT,
                           xaxis_title=_plot_title("Observation Index"),
                           yaxis_title="Residuals",
                           title="Residuals vs Observation Order")
        st.plotly_chart(fig2, use_container_width=True)

        st.markdown("**📖 Interpretation:**")
        if dw_ok:
            st.success(
                f"Durbin-Watson = {r['dw']:.3f} (ideal range: 1.5 – 2.5). No significant "
                "autocorrelation detected. Residuals appear independent of each other — "
                "a key assumption of OLS regression."
            )
        elif r["dw"] < 1.5:
            st.warning(
                f"DW = {r['dw']:.3f} < 1.5 — positive autocorrelation suspected. Consecutive "
                "residuals tend to be similar in sign. This is common in time-series data. "
                "Consider: (1) adding lag terms, (2) using time-series models (ARIMA), or "
                "(3) applying Newey-West standard errors."
            )
        else:
            st.warning(
                f"DW = {r['dw']:.3f} > 2.5 — negative autocorrelation suspected. "
                "Consecutive residuals tend to alternate in sign. Review data ordering "
                "and consider whether an alternating pattern exists in your data collection."
            )

    # ── 3. Homoscedasticity ───────────────────────────────
    with st.expander("📏 Test 3: Homoscedasticity  —  Breusch-Pagan", expanded=True):
        bp_ok = r["bp_p"] > alpha
        st.markdown(
            f"**Breusch-Pagan LM statistic:** `{r['bp_stat']:.3f}`,  "
            f"p = `{fmt_p(r['bp_p'])}`  →  {verdict_md(bp_ok)}"
        )
        abs_std = [abs(v) for v in r["std_resid"]]
        fig3 = go.Figure()
        fig3.add_trace(go.Scatter(
            x=r["fitted"], y=abs_std, mode="markers",
            marker=dict(color="#d97706", opacity=0.6, size=7),
            name="|Std. Residual|"
        ))
        fig3.update_layout(**PLOTLY_LIGHT,
                           xaxis_title=_plot_title("Fitted Values"),
                           yaxis_title="|Standardised Residuals|",
                           title="Scale-Location Plot (Spread vs Fitted)")
        st.plotly_chart(fig3, use_container_width=True)

        st.markdown("**📖 Interpretation:**")
        if bp_ok:
            st.success(
                f"Breusch-Pagan test is non-significant (p = {fmt_p(r['bp_p'])}). "
                "No evidence of heteroscedasticity — error variance appears constant across "
                "fitted values. The spread of residuals is roughly uniform, supporting OLS "
                "efficiency assumptions."
            )
        else:
            st.warning(
                f"Breusch-Pagan test is significant (p = {fmt_p(r['bp_p'])}) — "
                "heteroscedasticity detected. The spread of residuals changes with fitted values. "
                "This makes standard errors unreliable. Consider: (1) using HC3 robust standard "
                "errors, (2) log-transforming the dependent variable, or (3) Weighted Least Squares."
            )

    # ── 4. Normality ──────────────────────────────────────
    with st.expander("🔔 Test 4: Normality of Residuals  —  Shapiro-Wilk + Q-Q Plot", expanded=True):
        sw_ok = r["sw_p"] > alpha
        st.markdown(
            f"**Shapiro-Wilk W:** `{r['sw_stat']:.4f}`,  "
            f"p = `{fmt_p(r['sw_p'])}`  →  {verdict_md(sw_ok)}"
        )
        sorted_resid = np.sort(r["resid"])
        n_obs = len(sorted_resid)
        probs = (np.arange(1, n_obs+1) - 0.5) / n_obs
        theoretical = stats.norm.ppf(probs)

        fig4 = go.Figure()
        fig4.add_trace(go.Scatter(
            x=theoretical, y=sorted_resid, mode="markers",
            marker=dict(color="#7c3aed", opacity=0.6, size=7),
            name="Sample quantiles"
        ))
        mn, mx = theoretical[0], theoretical[-1]
        fig4.add_trace(go.Scatter(
            x=[mn, mx], y=[mn, mx], mode="lines",
            line=dict(color="#e53e3e", dash="dash", width=2),
            name="Normal reference line"
        ))
        fig4.update_layout(**PLOTLY_LIGHT,
                           xaxis_title=_plot_title("Theoretical Quantiles (Normal distribution)"),
                           yaxis_title="Sample Quantiles (Residuals)",
                           title="Normal Q-Q Plot of Residuals")
        st.plotly_chart(fig4, use_container_width=True)

        st.markdown("**📖 Interpretation:**")
        if sw_ok:
            st.success(
                f"Shapiro-Wilk test non-significant (p = {fmt_p(r['sw_p'])}). Residuals are "
                "approximately normally distributed. In the Q-Q plot, points fall close to the "
                "diagonal reference line. OLS t-tests and F-tests are valid under this condition."
            )
        else:
            st.warning(
                f"Shapiro-Wilk significant (p = {fmt_p(r['sw_p'])}) — residuals deviate from "
                "normality. In the Q-Q plot, look for S-curves (skewness) or heavy tails "
                "(kurtosis). With large samples (N > 100), OLS is often robust to mild "
                "non-normality (Central Limit Theorem). For small samples: consider "
                "transforming the DV or using bootstrap confidence intervals."
            )

    st.markdown("<br>", unsafe_allow_html=True)
    bc1, bc2 = st.columns([1, 5])
    with bc1:
        if st.button("‹ Back"):
            go_step(4); st.rerun()
    with bc2:
        if st.button("Diagnostics ›", type="primary"):
            go_step(6); st.rerun()


# ═══════════════════════════════════════════════════════════════════════
#  STEP 6 — DIAGNOSTICS + DOWNLOAD
# ═══════════════════════════════════════════════════════════════════════
elif cs == 6:
    st.markdown(badge("Step 6"), unsafe_allow_html=True)
    st.markdown("## Model Diagnostics")

    r     = st.session_state["model_result"]
    alpha = r["alpha"]
    d     = st.session_state["decimals"]

    if r is None:
        st.warning("Run the model first.")
        if st.button("‹ Back"):
            go_step(3); st.rerun()
        st.stop()

    n            = r["n"]
    resid_thresh = 2.5
    lev_thresh   = 2 * (len(st.session_state["selected_ivs"]) + 1) / n
    cook_thresh  = 1.0

    st.info(
        f"**Flagging thresholds:** |Standardised residual| > {resid_thresh} → Outlier  ·  "
        f"Leverage > {lev_thresh:.4f} → High leverage  ·  Cook's D > {cook_thresh} → Influential"
    )

    std_resid = np.array(r["std_resid"])
    leverage  = np.array(r["leverage"])
    cooks_d   = np.array(r["cooks_d"])
    obs       = np.array(r["obs_idx"])

    def obs_color(i):
        if cooks_d[i] > cook_thresh:         return "#e53e3e"
        if abs(std_resid[i]) > resid_thresh: return "#d97706"
        if leverage[i] > lev_thresh:         return "#7c3aed"
        return "#1a6bb5"

    colors = [obs_color(i) for i in obs]

    # ── Std Residuals ─────────────────────────────────────
    with st.expander("📍 Plot 1: Standardised Residuals", expanded=True):
        fig_sr = go.Figure()
        fig_sr.add_trace(go.Scatter(
            x=obs+1, y=std_resid, mode="markers",
            marker=dict(color=colors, size=8, opacity=0.8),
            text=[f"Obs {i+1}" for i in obs],
            hovertemplate="%{text}<br>Std residual: %{y:.3f}<extra></extra>"
        ))
        fig_sr.add_hline(y=resid_thresh,  line_dash="dash", line_color="#e53e3e", line_width=2,
                         annotation_text=f"+{resid_thresh}", annotation_position="top right")
        fig_sr.add_hline(y=-resid_thresh, line_dash="dash", line_color="#e53e3e", line_width=2,
                         annotation_text=f"-{resid_thresh}", annotation_position="bottom right")
        fig_sr.add_hline(y=0, line_dash="dot", line_color="#059669", line_width=1.5)
        fig_sr.update_layout(**PLOTLY_LIGHT, height=380,
                             xaxis_title=_plot_title("Observation Number"),
                             yaxis_title="Standardised Residual",
                             title="Standardised Residuals by Observation")
        st.plotly_chart(fig_sr, use_container_width=True)
        st.markdown("**📖 Interpretation:** Points beyond the dashed red lines (±2.5) are "
                    "potential **outliers** — their observed values differ substantially from "
                    "what the model predicts. Orange points exceed the threshold. Investigate "
                    "these observations for data entry errors or genuinely unusual cases.")

    # ── Leverage ─────────────────────────────────────────
    with st.expander("📍 Plot 2: Leverage (Hat Values)", expanded=True):
        lev_colors = ["#7c3aed" if leverage[i] > lev_thresh else "#1a6bb5" for i in obs]
        fig_lev = go.Figure()
        fig_lev.add_trace(go.Scatter(
            x=obs+1, y=leverage, mode="markers",
            marker=dict(color=lev_colors, size=8, opacity=0.8),
            text=[f"Obs {i+1}" for i in obs],
            hovertemplate="%{text}<br>Leverage: %{y:.4f}<extra></extra>"
        ))
        fig_lev.add_hline(y=lev_thresh, line_dash="dash", line_color="#7c3aed", line_width=2,
                          annotation_text=f"Threshold: {lev_thresh:.3f}",
                          annotation_position="top right")
        fig_lev.update_layout(**PLOTLY_LIGHT, height=380,
                              xaxis_title=_plot_title("Observation Number"),
                              yaxis_title="Leverage (Hat Value)",
                              title="Leverage by Observation")
        st.plotly_chart(fig_lev, use_container_width=True)
        st.markdown("**📖 Interpretation:** Leverage measures how far each observation's "
                    "**predictor values** are from the average. High-leverage points (purple, "
                    f"above threshold {lev_thresh:.3f} = 2(k+1)/n) have unusual X-values and "
                    "can disproportionately pull the regression line. High leverage alone is "
                    "not harmful — it becomes a problem only if combined with a large residual.")

    # ── Cook's Distance ───────────────────────────────────
    with st.expander("📍 Plot 3: Cook's Distance", expanded=True):
        cook_colors = ["#e53e3e" if cooks_d[i] > cook_thresh else "#1a6bb5" for i in obs]
        fig_cd = go.Figure()
        fig_cd.add_trace(go.Bar(
            x=obs+1, y=cooks_d, marker_color=cook_colors, opacity=0.85,
            text=[f"Obs {i+1}" for i in obs],
            hovertemplate="%{text}<br>Cook's D: %{y:.4f}<extra></extra>"
        ))
        fig_cd.add_hline(y=cook_thresh, line_dash="dash", line_color="#e53e3e", line_width=2,
                         annotation_text="D = 1 (threshold)",
                         annotation_position="top right")
        fig_cd.update_layout(**PLOTLY_LIGHT, height=380,
                             xaxis_title=_plot_title("Observation Number"),
                             yaxis_title="Cook's Distance",
                             title="Cook's Distance — Influence of Each Observation")
        st.plotly_chart(fig_cd, use_container_width=True)
        st.markdown("**📖 Interpretation:** Cook's Distance combines leverage and residual "
                    "size into a single influence measure. It estimates how much all fitted "
                    "values would change if that observation were removed. Values **> 1** "
                    "(red bars) indicate highly influential observations that substantially "
                    "affect the model coefficients — these warrant careful scrutiny.")

    # ── Influence Map ─────────────────────────────────────
    with st.expander("📍 Plot 4: Influence Map  (Leverage × Cook's D)", expanded=True):
        sizes = [max(9, min(35, abs(std_resid[i]) * 7)) for i in obs]
        fig_inf = go.Figure()
        fig_inf.add_trace(go.Scatter(
            x=leverage, y=cooks_d, mode="markers",
            marker=dict(color=colors, size=sizes, opacity=0.75,
                        line=dict(width=1, color="#ffffff")),
            text=[f"Obs {i+1}<br>Std resid: {std_resid[i]:.3f}<br>"
                  f"Leverage: {leverage[i]:.4f}<br>Cook's D: {cooks_d[i]:.4f}"
                  for i in obs],
            hovertemplate="%{text}<extra></extra>"
        ))
        fig_inf.add_vline(x=lev_thresh, line_dash="dash", line_color="#7c3aed", line_width=1.5,
                          annotation_text="High leverage", annotation_position="top left")
        fig_inf.add_hline(y=cook_thresh, line_dash="dash", line_color="#e53e3e", line_width=1.5,
                          annotation_text="High influence", annotation_position="top right")
        fig_inf.update_layout(**PLOTLY_LIGHT, height=430,
                              xaxis_title=_plot_title("Leverage (Hat Value)"),
                              yaxis_title="Cook's Distance",
                              title="Influence Map  —  Bubble size = |Std. Residual|")
        st.plotly_chart(fig_inf, use_container_width=True)
        st.markdown("**📖 Interpretation:** This chart combines all three diagnostic dimensions. "
                    "Points in the **upper-right quadrant** (high leverage AND high Cook's D) "
                    "are most problematic — they have unusual predictor values AND strongly "
                    "influence the results. Bubble size reflects the magnitude of the standardised "
                    "residual. Colour: 🔵 Normal · 🟠 Outlier · 🟣 High leverage · 🔴 Influential.")

    # ── Flagged observations table ────────────────────────
    st.markdown("---")
    st.markdown("### Flagged Observations Summary")
    flagged_idx = [i for i in obs
                   if abs(std_resid[i]) > resid_thresh
                   or leverage[i] > lev_thresh
                   or cooks_d[i] > cook_thresh]

    if not flagged_idx:
        st.success("✅ No observations flagged under the current thresholds. "
                   "The model appears free of major outliers, high-leverage, or influential points.")
    else:
        flag_rows = []
        for i in flagged_idx:
            flags = []
            if abs(std_resid[i]) > resid_thresh: flags.append("Outlier")
            if leverage[i] > lev_thresh:         flags.append("High leverage")
            if cooks_d[i] > cook_thresh:         flags.append("Influential")
            flag_rows.append({
                "Obs #":           i+1,
                "Std. Residual":   f"{std_resid[i]:.3f}",
                "Leverage":        f"{leverage[i]:.4f}",
                "Cook's D":        f"{cooks_d[i]:.4f}",
                "Flags":           ", ".join(flags),
            })
        flag_df = pd.DataFrame(flag_rows)
        st.dataframe(flag_df, use_container_width=True, hide_index=True)
        st.markdown(
            f"**{len(flagged_idx)} observation(s) flagged.** Review these carefully — "
            "they may reflect genuine extreme values, data errors, or important subgroups."
        )

    st.markdown("<br>", unsafe_allow_html=True)
    bc1, bc2 = st.columns([1, 5])
    with bc1:
        if st.button("‹ Back"):
            go_step(5); st.rerun()
    with bc2:
        if st.button("Multicollinearity ›", type="primary"):
            go_step(7); st.rerun()


# ═══════════════════════════════════════════════════════════════════════
#  STEP 7 — MULTICOLLINEARITY + DOWNLOAD
# ═══════════════════════════════════════════════════════════════════════
elif cs == 7:
    st.markdown(badge("Step 7"), unsafe_allow_html=True)
    st.markdown("## Multicollinearity Assessment")

    r     = st.session_state["model_result"]
    alpha = r["alpha"]
    d     = st.session_state["decimals"]

    if r is None:
        st.warning("Run the model first.")
        if st.button("‹ Back"):
            go_step(3); st.rerun()
        st.stop()

    # ── VIF Table ─────────────────────────────────────────
    st.markdown("### Variance Inflation Factor (VIF)")
    st.markdown(
        "VIF quantifies how much the variance of a coefficient is inflated due to "
        "correlation with other predictors. **VIF = 1** means no collinearity. "
        "**VIF > 5** warrants attention; **VIF > 10** indicates serious multicollinearity."
    )

    vif_rows = r.get("vif_rows", [])

    if not vif_rows:
        st.info("VIF requires at least 2 numeric predictors in the model.")
    else:
        vif_display = []
        for vr in vif_rows:
            vif_val = vr["vif"]
            if np.isnan(vif_val):
                badge_html = '<span class="vif-mod">N/A</span>'
                concern    = "Cannot compute"
            elif vif_val < 5:
                badge_html = f'<span class="vif-ok">✅ Low ({vif_val:.2f})</span>'
                concern    = "No concern"
            elif vif_val < 10:
                badge_html = f'<span class="vif-mod">⚠️ Moderate ({vif_val:.2f})</span>'
                concern    = "Moderate — monitor"
            else:
                badge_html = f'<span class="vif-high">🔴 High ({vif_val:.2f})</span>'
                concern    = "Serious — action recommended"
            vif_display.append({
                "Predictor Term": vr["term"],
                "VIF":            f"{vif_val:.3f}" if not np.isnan(vif_val) else "N/A",
                "Assessment":     concern,
            })

        vif_df = pd.DataFrame(vif_display)

        def vif_highlight(row):
            try:
                v = float(row["VIF"])
            except Exception:
                return [""] * len(row)
            if v >= 10:
                return ["background-color: #f8d7da"] * len(row)
            elif v >= 5:
                return ["background-color: #fff3cd"] * len(row)
            else:
                return ["background-color: #d4edda"] * len(row)

        styled_vif = (
            vif_df.style
            .apply(vif_highlight, axis=1)
            .set_properties(**{"font-size": "14px", "padding": "9px 12px"})
            .set_table_styles([
                {"selector": "th", "props": [
                    ("background-color", "#1a3a6b"), ("color", "#ffffff"),
                    ("font-size", "13px"), ("font-weight", "700"), ("padding", "10px 12px"),
                ]},
            ])
        )
        st.dataframe(styled_vif, use_container_width=True, hide_index=True)

        # VIF chart
        vif_vals  = [vr["vif"] if not np.isnan(vr["vif"]) else 0 for vr in vif_rows]
        vif_names = [vr["term"] for vr in vif_rows]
        vif_bar_colors = ["#e53e3e" if v >= 10 else "#d97706" if v >= 5 else "#059669"
                          for v in vif_vals]

        fig_vif = go.Figure()
        fig_vif.add_trace(go.Bar(
            x=vif_names, y=vif_vals,
            marker_color=vif_bar_colors, opacity=0.85,
            text=[f"{v:.2f}" for v in vif_vals],
            textposition="outside",
            textfont=dict(size=13, color="#1a2340")
        ))
        fig_vif.add_hline(y=5,  line_dash="dash", line_color="#d97706", line_width=2,
                          annotation_text="VIF = 5 (caution)", annotation_position="top right")
        fig_vif.add_hline(y=10, line_dash="dash", line_color="#e53e3e", line_width=2,
                          annotation_text="VIF = 10 (serious)", annotation_position="top right")
        fig_vif.update_layout(**PLOTLY_LIGHT, height=400,
                              xaxis_title=_plot_title("Predictor"),
                              yaxis_title="Variance Inflation Factor (VIF)",
                              title="VIF by Predictor  —  Multicollinearity Assessment")
        st.plotly_chart(fig_vif, use_container_width=True)

        st.markdown("**📖 Interpretation:**")
        high_vif = [vr["term"] for vr in vif_rows if not np.isnan(vr["vif"]) and vr["vif"] >= 10]
        mod_vif  = [vr["term"] for vr in vif_rows if not np.isnan(vr["vif"]) and 5 <= vr["vif"] < 10]
        if not high_vif and not mod_vif:
            st.success(
                "All VIF values are below 5. No meaningful multicollinearity detected. "
                "Your predictors are sufficiently independent — coefficient estimates are stable "
                "and standard errors are not inflated."
            )
        else:
            msg = []
            if high_vif:
                msg.append(
                    f"**Serious multicollinearity** detected for: {', '.join(high_vif)} (VIF ≥ 10). "
                    "These predictors are highly correlated with others. Coefficients may be "
                    "unstable and standard errors inflated. Consider: removing one of the "
                    "correlated predictors, combining them into a composite score, or using "
                    "ridge regression."
                )
            if mod_vif:
                msg.append(
                    f"**Moderate multicollinearity** for: {', '.join(mod_vif)} (VIF 5–10). "
                    "Monitor these, especially if coefficient signs are unexpected or SEs are large."
                )
            for m in msg:
                st.warning(m)

    st.markdown("---")

    # ── Correlation Heatmap ───────────────────────────────
    st.markdown("### Correlation Matrix  (Numeric Predictors)")
    num_ivs   = r.get("num_ivs", [])
    corr_dict = r.get("corr_matrix", None)

    if corr_dict and len(num_ivs) >= 2:
        corr_df = pd.DataFrame(corr_dict)

        z    = corr_df.values.tolist()
        text = [[f"{val:.3f}" for val in row] for row in z]

        fig_corr = go.Figure(go.Heatmap(
            z=z,
            x=corr_df.columns.tolist(),
            y=corr_df.index.tolist(),
            text=text,
            texttemplate="%{text}",
            textfont=dict(size=13, color="#1a2340"),
            colorscale="RdBu",
            zmid=0,
            zmin=-1, zmax=1,
            colorbar=dict(
                title=dict(text="Pearson r", font=dict(size=13)),
                tickfont=dict(size=12),
            ),
            showscale=True,
        ))
        fig_corr.update_layout(
            **PLOTLY_LIGHT_HEATMAP,
            height=max(380, 80 * len(num_ivs)),
            title=_plot_title("Pearson Correlation Matrix of Numeric Predictors"),
        )
        st.plotly_chart(fig_corr, use_container_width=True)

        st.markdown("**📖 Interpretation:** Values close to **+1 or −1** indicate strong "
                    "correlation between two predictors, which contributes to multicollinearity. "
                    "Values near **0** indicate independence. As a rule of thumb, |r| > 0.7 "
                    "between any two predictors warrants attention. The correlation matrix "
                    "complements VIF by showing which specific pairs of predictors are correlated.")
    else:
        st.info("Correlation matrix requires at least 2 numeric predictors.")

    # ═══════════════════════════════════════════════════════
    #  WORD REPORT DOWNLOAD
    # ═══════════════════════════════════════════════════════
    st.markdown("---")
    st.markdown("### 📥 Download Full Report (.docx)")
    st.markdown(
        "The report includes all tables (coefficients, fit metrics, assumption tests, "
        "diagnostics, VIF) and all charts embedded as images."
    )

    # ── docx helpers ──────────────────────────────────────
    def _set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_col_width(table, col_idx, width_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width_cm)

    def _add_heading(doc, text, level=1):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = (
            RGBColor(0x1A, 0x3A, 0x6B) if level == 1 else RGBColor(0x2C, 0x4A, 0x7C)
        )
        h.runs[0].font.size = Pt(16 if level == 1 else 13)
        h.runs[0].font.bold = True
        return h

    def _add_note(doc, text, italic=True):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(11)
        p.runs[0].italic = italic
        p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return p

    def _fig_to_img_bytes(fig, width=900, height=420):
        """Convert plotly figure to PNG bytes for embedding."""
        try:
            return fig.to_image(format="png", width=width, height=height, scale=2)
        except Exception:
            return None

    def generate_docx_report() -> bytes:
        r2    = st.session_state["model_result"]
        d2    = st.session_state["decimals"]
        alpha2 = r2["alpha"]
        ci_pct = int(r2["ci_level"] * 100)

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.2)
            section.bottom_margin = Cm(2.2)
            section.left_margin   = Cm(2.8)
            section.right_margin  = Cm(2.8)

        # Default style
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # ── Title page ─────────────────────────────────────
        title = doc.add_heading("OLS Regression Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(22)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(0x1A, 0x2A, 0x50)

        sub = doc.add_paragraph("Ordinary Least Squares — Full Output")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(13)
        sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()

        # ── Section 1: Model Specification ─────────────────
        _add_heading(doc, "1. Model Specification")
        spec_items = [
            ("Dependent variable",     r2["dv"]),
            ("Independent variables",  ", ".join(st.session_state["selected_ivs"])),
            ("Patsy formula",          r2["formula"]),
            ("Significance level (α)", str(alpha2)),
            ("Confidence interval",    f"{ci_pct}%"),
            ("Mean-centered vars",     ", ".join([k for k, v in
                                        st.session_state["centering"].items() if v]) or "None"),
            ("Observations (N)",       str(r2["n"])),
        ]
        tbl = doc.add_table(rows=len(spec_items), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(spec_items):
            cell0 = tbl.rows[i].cells[0]
            cell1 = tbl.rows[i].cells[1]
            cell0.text = k
            cell1.text = str(v)
            cell0.paragraphs[0].runs[0].bold = True
            cell0.paragraphs[0].runs[0].font.size = Pt(11)
            cell1.paragraphs[0].runs[0].font.size = Pt(11)
            _set_cell_bg(cell0, "D9E2F3")
            if i % 2 == 1:
                _set_cell_bg(cell1, "F5F7FA")
        doc.add_paragraph()

        # ── Section 2: Coefficient Table ───────────────────
        _add_heading(doc, "2. Coefficient Table")
        _add_note(doc,
            f"α = {alpha2}  ·  CI level = {ci_pct}%  ·  "
            "Green shading = significant  ·  Red shading = not significant"
        )

        ci_lo_lbl = f"CI {(1-r2['ci_level'])/2*100:.1f}%"
        ci_hi_lbl = f"CI {(1-(1-r2['ci_level'])/2)*100:.1f}%"
        hdr = ["Term", "Coefficient", "Std. Error", "t-statistic", "p-value",
               ci_lo_lbl, ci_hi_lbl, "Significant"]

        ctbl = doc.add_table(rows=1 + len(r2["coef_rows"]), cols=len(hdr))
        ctbl.style = "Table Grid"

        for j, h in enumerate(hdr):
            cell = ctbl.rows[0].cells[j]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1A3A6B")

        for i, row in enumerate(r2["coef_rows"], start=1):
            vals = [
                row["term"],
                f"{row['coef']:.{d2}f}",
                f"{row['se']:.{d2}f}",
                f"{row['tstat']:.{d2}f}",
                fmt_p(row["pval"]),
                f"{row['ci_lo']:.{d2}f}",
                f"{row['ci_hi']:.{d2}f}",
                "Yes ***" if row["sig"] else "No",
            ]
            bg = ("C6EFCE" if row["sig"]
                  else ("FFC7CE" if row["term"] != "Intercept" else "F2F2F2"))
            for j, v in enumerate(vals):
                cell = ctbl.rows[i].cells[j]
                cell.text = v
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(11)
                run.font.name = "Arial"
                _set_cell_bg(cell, bg)

        doc.add_paragraph()

        # F-test row
        fp = doc.add_paragraph()
        r_run = fp.add_run(
            f"F({r2['df_model']}, {r2['df_resid']}) = {r2['fstat']:.3f}  |  "
            f"p = {fmt_p(r2['fpval'])}  |  R² = {r2['rsq']:.4f}  |  "
            f"Adj. R² = {r2['rsq_adj']:.4f}  |  RMSE = {r2['rmse']:.{d2}f}  |  "
            f"AIC = {r2['aic']:.1f}  |  BIC = {r2['bic']:.1f}"
        )
        r_run.font.size  = Pt(11)
        r_run.font.bold  = True
        r_run.font.name  = "Arial"
        doc.add_paragraph()

        # ── Section 3: Fit Metrics & Assumptions ───────────
        _add_heading(doc, "3. Fit Metrics & Assumption Tests")

        fit_items = [
            ("R²",           f"{r2['rsq']:.4f}",     f"Model explains {r2['rsq']*100:.1f}% of variance"),
            ("Adjusted R²",  f"{r2['rsq_adj']:.4f}", "Penalised for number of predictors"),
            ("RMSE",         f"{r2['rmse']:.{d2}f}", "Root Mean Square Error"),
            ("AIC",          f"{r2['aic']:.2f}",      "Lower = better model fit"),
            ("BIC",          f"{r2['bic']:.2f}",      "Stricter penalisation than AIC"),
        ]
        ftbl = doc.add_table(rows=1 + len(fit_items), cols=3)
        ftbl.style = "Table Grid"
        for j, h in enumerate(["Metric", "Value", "Note"]):
            cell = ftbl.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1A3A6B")
        for i, (metric, val, note) in enumerate(fit_items, start=1):
            cells = ftbl.rows[i].cells
            cells[0].text = metric
            cells[1].text = val
            cells[2].text = note
            for c in cells:
                c.paragraphs[0].runs[0].font.size = Pt(11)
            _set_cell_bg(cells[0], "D9E2F3")
            if i % 2 == 1:
                _set_cell_bg(cells[1], "F5F7FA")
        doc.add_paragraph()

        # Assumption results
        _add_heading(doc, "Assumption Test Results", level=2)
        lin_ok = abs(r2["lin_r"]) < 0.1
        dw_ok  = 1.5 < r2["dw"] < 2.5
        bp_ok  = r2["bp_p"] > alpha2
        sw_ok  = r2["sw_p"] > alpha2

        assume_items = [
            ("Linearity (Pearson r resid~fitted)",
             f"r = {r2['lin_r']:.4f}, p = {fmt_p(r2['lin_p'])}", lin_ok),
            ("Independence (Durbin-Watson)",
             f"DW = {r2['dw']:.4f}", dw_ok),
            ("Homoscedasticity (Breusch-Pagan)",
             f"LM = {r2['bp_stat']:.3f}, p = {fmt_p(r2['bp_p'])}", bp_ok),
            ("Normality of Residuals (Shapiro-Wilk)",
             f"W = {r2['sw_stat']:.4f}, p = {fmt_p(r2['sw_p'])}", sw_ok),
        ]
        atbl = doc.add_table(rows=1 + len(assume_items), cols=3)
        atbl.style = "Table Grid"
        for j, h in enumerate(["Assumption", "Statistic", "Verdict"]):
            cell = atbl.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1A3A6B")
        for i, (name, stat_str, ok) in enumerate(assume_items, start=1):
            cells = atbl.rows[i].cells
            cells[0].text = name
            cells[1].text = stat_str
            cells[2].text = "✓ Supported" if ok else "⚠ Caution"
            for c in cells:
                c.paragraphs[0].runs[0].font.size = Pt(11)
            _set_cell_bg(cells[2], "C6EFCE" if ok else "FFC7CE")
        doc.add_paragraph()

        # ── Assumption Plots ────────────────────────────────
        _add_heading(doc, "Assumption Diagnostic Plots", level=2)

        # Residuals vs Fitted
        fig_rf = go.Figure()
        fig_rf.add_trace(go.Scatter(x=r2["fitted"], y=r2["resid"], mode="markers",
                                    marker=dict(color="#1a6bb5", opacity=0.6, size=7)))
        fig_rf.add_hline(y=0, line_dash="dash", line_color="#e53e3e", line_width=2)
        fig_rf.update_layout(**PLOTLY_LIGHT, xaxis_title=_plot_title("Fitted Values"),
                             yaxis_title="Residuals", title="Residuals vs Fitted Values",
                             height=380)
        img1 = _fig_to_img_bytes(fig_rf)
        if img1:
            doc.add_picture(io.BytesIO(img1), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_note(doc, "Residuals vs Fitted: Random scatter around zero line supports linearity.")
        doc.add_paragraph()

        # Q-Q Plot
        sorted_resid = np.sort(r2["resid"])
        n_obs        = len(sorted_resid)
        probs_arr    = (np.arange(1, n_obs+1) - 0.5) / n_obs
        theoretical  = stats.norm.ppf(probs_arr)
        mn2, mx2     = theoretical[0], theoretical[-1]
        fig_qq = go.Figure()
        fig_qq.add_trace(go.Scatter(x=theoretical, y=sorted_resid, mode="markers",
                                    marker=dict(color="#7c3aed", opacity=0.6, size=7)))
        fig_qq.add_trace(go.Scatter(x=[mn2, mx2], y=[mn2, mx2], mode="lines",
                                    line=dict(color="#e53e3e", dash="dash", width=2)))
        fig_qq.update_layout(**PLOTLY_LIGHT,
                             xaxis_title=_plot_title("Theoretical Quantiles"),
                             yaxis_title="Sample Quantiles",
                             title="Normal Q-Q Plot of Residuals", height=380)
        img2 = _fig_to_img_bytes(fig_qq)
        if img2:
            doc.add_picture(io.BytesIO(img2), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_note(doc, "Q-Q Plot: Points on the diagonal indicate normally distributed residuals.")
        doc.add_paragraph()

        # Scale-Location
        abs_std2 = [abs(v) for v in r2["std_resid"]]
        fig_sl = go.Figure()
        fig_sl.add_trace(go.Scatter(x=r2["fitted"], y=abs_std2, mode="markers",
                                    marker=dict(color="#d97706", opacity=0.6, size=7)))
        fig_sl.update_layout(**PLOTLY_LIGHT, xaxis_title=_plot_title("Fitted Values"),
                             yaxis_title="|Standardised Residuals|",
                             title="Scale-Location Plot", height=380)
        img3 = _fig_to_img_bytes(fig_sl)
        if img3:
            doc.add_picture(io.BytesIO(img3), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_note(doc, "Scale-Location: Even spread across fitted values supports homoscedasticity.")
        doc.add_paragraph()

        # ── Section 4: Diagnostics ─────────────────────────
        _add_heading(doc, "4. Model Diagnostics")
        n2          = r2["n"]
        lev_thresh2 = 2 * (len(st.session_state["selected_ivs"]) + 1) / n2
        cook_thresh2 = 1.0
        resid_thresh2 = 2.5

        std_resid2 = np.array(r2["std_resid"])
        leverage2  = np.array(r2["leverage"])
        cooks_d2   = np.array(r2["cooks_d"])
        obs2       = np.array(r2["obs_idx"])

        _add_note(doc,
            f"Thresholds: |Std. residual| > {resid_thresh2} = Outlier  ·  "
            f"Leverage > {lev_thresh2:.4f} = High leverage  ·  Cook's D > {cook_thresh2} = Influential"
        )
        doc.add_paragraph()

        # Std residuals plot
        def obs_color2(i):
            if cooks_d2[i] > cook_thresh2:          return "#e53e3e"
            if abs(std_resid2[i]) > resid_thresh2:  return "#d97706"
            if leverage2[i] > lev_thresh2:          return "#7c3aed"
            return "#1a6bb5"
        colors2 = [obs_color2(i) for i in obs2]

        fig_sr2 = go.Figure()
        fig_sr2.add_trace(go.Scatter(x=obs2+1, y=std_resid2, mode="markers",
                                     marker=dict(color=colors2, size=8, opacity=0.8)))
        fig_sr2.add_hline(y=resid_thresh2,  line_dash="dash", line_color="#e53e3e", line_width=2)
        fig_sr2.add_hline(y=-resid_thresh2, line_dash="dash", line_color="#e53e3e", line_width=2)
        fig_sr2.add_hline(y=0, line_dash="dot", line_color="#059669", line_width=1.5)
        fig_sr2.update_layout(**PLOTLY_LIGHT, height=380,
                              xaxis_title=_plot_title("Observation Number"),
                              yaxis_title="Standardised Residual",
                              title="Standardised Residuals")
        img4 = _fig_to_img_bytes(fig_sr2)
        if img4:
            doc.add_picture(io.BytesIO(img4), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_note(doc, "Standardised Residuals: Points beyond ±2.5 are potential outliers.")
        doc.add_paragraph()

        # Cook's distance
        cook_colors2 = ["#e53e3e" if cooks_d2[i] > cook_thresh2 else "#1a6bb5" for i in obs2]
        fig_cd2 = go.Figure()
        fig_cd2.add_trace(go.Bar(x=obs2+1, y=cooks_d2, marker_color=cook_colors2, opacity=0.85))
        fig_cd2.add_hline(y=cook_thresh2, line_dash="dash", line_color="#e53e3e", line_width=2)
        fig_cd2.update_layout(**PLOTLY_LIGHT, height=380,
                              xaxis_title=_plot_title("Observation Number"),
                              yaxis_title="Cook's Distance",
                              title="Cook's Distance — Influential Observations")
        img5 = _fig_to_img_bytes(fig_cd2)
        if img5:
            doc.add_picture(io.BytesIO(img5), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_note(doc, "Cook's Distance: Values > 1 indicate highly influential observations.")
        doc.add_paragraph()

        # Flagged table
        flagged2 = [i for i in obs2
                    if abs(std_resid2[i]) > resid_thresh2
                    or leverage2[i] > lev_thresh2
                    or cooks_d2[i] > cook_thresh2]

        if not flagged2:
            doc.add_paragraph("✓ No flagged observations under the current thresholds.").runs[0].font.size = Pt(11)
        else:
            dhdr = ["Obs #", "Std. Residual", "Leverage", "Cook's D", "Flags"]
            dtbl = doc.add_table(rows=1 + len(flagged2), cols=5)
            dtbl.style = "Table Grid"
            for j, h in enumerate(dhdr):
                cell = dtbl.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, "1A3A6B")
            for row_num, i in enumerate(flagged2, start=1):
                flags2 = []
                if abs(std_resid2[i]) > resid_thresh2: flags2.append("Outlier")
                if leverage2[i] > lev_thresh2:          flags2.append("High leverage")
                if cooks_d2[i] > cook_thresh2:          flags2.append("Influential")
                vals2 = [str(i+1), f"{std_resid2[i]:.3f}", f"{leverage2[i]:.4f}",
                         f"{cooks_d2[i]:.4f}", ", ".join(flags2)]
                for j, v in enumerate(vals2):
                    cell = dtbl.rows[row_num].cells[j]
                    cell.text = v
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    _set_cell_bg(cell, "FFC7CE")
        doc.add_paragraph()

        # ── Section 5: Multicollinearity ───────────────────
        _add_heading(doc, "5. Multicollinearity Assessment")

        vif_rows2 = r2.get("vif_rows", [])
        if not vif_rows2:
            doc.add_paragraph("VIF requires at least 2 numeric predictors.").runs[0].font.size = Pt(11)
        else:
            vhdr = ["Predictor Term", "VIF", "Assessment"]
            vtbl = doc.add_table(rows=1 + len(vif_rows2), cols=3)
            vtbl.style = "Table Grid"
            for j, h in enumerate(vhdr):
                cell = vtbl.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, "1A3A6B")
            for i, vr in enumerate(vif_rows2, start=1):
                vv = vr["vif"]
                vif_str = f"{vv:.3f}" if not np.isnan(vv) else "N/A"
                if np.isnan(vv):
                    concern2 = "Cannot compute"; bg_v = "F2F2F2"
                elif vv < 5:
                    concern2 = "No concern (< 5)"; bg_v = "C6EFCE"
                elif vv < 10:
                    concern2 = "Moderate (5–10)"; bg_v = "FFEB9C"
                else:
                    concern2 = "Serious (≥ 10)"; bg_v = "FFC7CE"
                vals3 = [vr["term"], vif_str, concern2]
                for j, v in enumerate(vals3):
                    cell = vtbl.rows[i].cells[j]
                    cell.text = v
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    if j == 2:
                        _set_cell_bg(cell, bg_v)
            doc.add_paragraph()

            # VIF chart in report
            vif_vals2  = [vr["vif"] if not np.isnan(vr["vif"]) else 0 for vr in vif_rows2]
            vif_names2 = [vr["term"] for vr in vif_rows2]
            vif_bar2   = ["#e53e3e" if v >= 10 else "#d97706" if v >= 5 else "#059669"
                          for v in vif_vals2]
            fig_vif2 = go.Figure()
            fig_vif2.add_trace(go.Bar(x=vif_names2, y=vif_vals2, marker_color=vif_bar2,
                                      opacity=0.85, text=[f"{v:.2f}" for v in vif_vals2],
                                      textposition="outside",
                                      textfont=dict(size=13)))
            fig_vif2.add_hline(y=5,  line_dash="dash", line_color="#d97706", line_width=2)
            fig_vif2.add_hline(y=10, line_dash="dash", line_color="#e53e3e", line_width=2)
            fig_vif2.update_layout(**PLOTLY_LIGHT, height=400,
                                   xaxis_title=_plot_title("Predictor"), yaxis_title="VIF",
                                   title="VIF by Predictor — Multicollinearity Assessment")
            img6 = _fig_to_img_bytes(fig_vif2)
            if img6:
                doc.add_picture(io.BytesIO(img6), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_note(doc, "VIF: Green < 5 (OK), Amber 5–10 (caution), Red ≥ 10 (serious).")
            doc.add_paragraph()

        # Correlation heatmap in report
        corr_dict2 = r2.get("corr_matrix", None)
        num_ivs2   = r2.get("num_ivs", [])
        if corr_dict2 and len(num_ivs2) >= 2:
            _add_heading(doc, "Predictor Correlation Matrix", level=2)
            corr_df2 = pd.DataFrame(corr_dict2)
            z2   = corr_df2.values.tolist()
            txt2 = [[f"{val:.3f}" for val in row] for row in z2]
            fig_corr2 = go.Figure(go.Heatmap(
                z=z2, x=corr_df2.columns.tolist(), y=corr_df2.index.tolist(),
                text=txt2, texttemplate="%{text}",
                textfont=dict(size=13),
                colorscale="RdBu", zmid=0, zmin=-1, zmax=1,
                showscale=True,
            ))
            fig_corr2.update_layout(**PLOTLY_LIGHT_HEATMAP,
                                    height=max(350, 80 * len(num_ivs2)),
                                    title=_plot_title("Pearson Correlation Matrix of Numeric Predictors"))
            img7 = _fig_to_img_bytes(fig_corr2, width=700,
                                     height=max(350, 80 * len(num_ivs2)))
            if img7:
                doc.add_picture(io.BytesIO(img7), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_note(doc, "Correlation matrix: |r| > 0.7 between predictors warrants attention.")
            doc.add_paragraph()

        # End
        end = doc.add_paragraph("— End of Report —")
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        end.runs[0].font.size = Pt(10)
        end.runs[0].italic = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ── Download button ───────────────────────────────────
    col_dl, _ = st.columns([2, 4])
    with col_dl:
        with st.spinner("Generating report — embedding charts…"):
            docx_bytes = generate_docx_report()
        st.download_button(
            label="⬇ Download Full Report (.docx)",
            data=docx_bytes,
            file_name="OLS_Regression_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("‹ Back to Diagnostics"):
        go_step(6); st.rerun()